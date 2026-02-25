import streamlit as st
import pandas as pd
from docx import Document
import io
import zipfile
import os

st.set_page_config(page_title="Генератор писем", layout="centered")

st.title("✉️ Автоматическое создание писем")
st.write("Загрузите Excel-таблицу и шаблон Word, чтобы сгенерировать письма.")

# 1. Загрузка файлов
uploaded_excel = st.file_uploader("Выберите Excel файл", type=["xlsx"])
uploaded_template = st.file_uploader("Выберите шаблон Word (.docx)", type=["docx"])

if uploaded_excel and uploaded_template:
    df = pd.read_excel(uploaded_excel)
    
    st.success("Файлы загружены!")
    
    if st.button("Сгенерировать письма"):
        # Буфер для создания ZIP-архива в памяти
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
            for index, row in df.iterrows():
                # Берем данные (индексы колонок как в вашем коде)
                obrashenie = str(row.iloc[10]) if not pd.isna(row.iloc[10]) else ""
                familiya = str(row.iloc[8]) if not pd.isna(row.iloc[8]) else f"Document_{index}"

                # Читаем шаблон из буфера
                uploaded_template.seek(0)
                doc = Document(uploaded_template)

                # Функция замены (параграфы и таблицы)
                def replace_text(doc_obj, old, new):
                    for p in doc_obj.paragraphs:
                        if old in p.text:
                            p.text = p.text.replace(old, new)
                    for table in doc_obj.tables:
                        for r in table.rows:
                            for cell in r.cells:
                                if old in cell.text:
                                    cell.text = cell.text.replace(old, new)

                replace_text(doc, "Обращения", obrashenie)
                replace_text(doc, "Фамилия", familiya)

                # Сохраняем готовый документ в память
                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                
                # Добавляем в ZIP
                zip_file.writestr(f"{familiya}.docx", doc_io.getvalue())

        # Кнопка скачивания архива
        st.download_button(
            label="Скачать все письма (ZIP)",
            data=zip_buffer.getvalue(),
            file_name="letters.zip",
            mime="application/zip"
        )
